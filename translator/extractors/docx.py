"""DOCX extractor using python-docx."""
from pathlib import Path
import re


def extract(docx_path: Path) -> str:
    """
    Extract text content from .docx file.

    Args:
        docx_path: Path to .docx file

    Returns:
        Extracted content as Markdown-formatted string
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: uv pip install python-docx")

    doc = Document(str(docx_path))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            parts.append('')
            continue

        # Check heading style
        style_name = para.style.name if para.style else ''
        if style_name.startswith('Heading'):
            try:
                level = int(re.search(r'\d+', style_name).group())
                parts.append(f"{'#' * level} {text}")
            except (ValueError, AttributeError):
                parts.append(f"## {text}")
        else:
            # Regular paragraph
            parts.append(text)

    # Extract tables
    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(' | '.join(cells))
        if rows_data:
            parts.append('')
            parts.append('| ' + ' | '.join(['---'] * len(table.columns)) + ' |')
            for row in rows_data:
                parts.append(f'| {row} |')

    return '\n'.join(parts)


def get_language(docx_path: Path) -> str:
    """Detect language from DOCX content."""
    content = extract(docx_path)
    # Use same detection as markdown
    cyrillic = len(re.findall(r'[\u0400-\u04FF]', content))
    latin = len(re.findall(r'[a-zA-Z]', content))
    total = len(content)
    if total == 0:
        return 'unknown'
    if cyrillic / total > 0.3:
        return 'ru'
    elif latin / total > 0.5:
        return 'en'
    return 'unknown'


# For compatibility with pipeline
def run(workspace: Path, **kwargs) -> dict:
    """Stub compatible with pipeline stage interface."""
    return {"status": "stub", "stage": "extract_docx"}